from docx import Document
import sys

def is_table_line(line):
    return '|' in line


def parse_table(block_lines):
    rows = []
    for l in block_lines:
        parts = [c.strip() for c in l.strip().strip('|').split('|')]
        rows.append(parts)
    # remove separator line if present (e.g., ---|---)
    if len(rows) >= 2 and all(set(c.replace('-','').strip())==set('') for c in rows[1]):
        rows.pop(1)
    return rows


def md_to_docx(md_path, docx_path):
    doc = Document()
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    n = len(lines)
    table_block = []
    while i < n:
        line = lines[i].rstrip('\n')
        if line.strip() == '':
            # flush table if any
            if table_block:
                rows = parse_table(table_block)
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Table Grid'
                    for r_idx, row in enumerate(rows):
                        for c_idx, cell in enumerate(row):
                            table.rows[r_idx].cells[c_idx].text = cell
                table_block = []
            else:
                doc.add_paragraph('')
            i += 1
            continue

        # heading
        if line.startswith('#'):
            if table_block:
                rows = parse_table(table_block)
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Table Grid'
                    for r_idx, row in enumerate(rows):
                        for c_idx, cell in enumerate(row):
                            table.rows[r_idx].cells[c_idx].text = cell
                table_block = []
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            if level >=1 and level <= 4:
                p = doc.add_paragraph(text)
                p.style = f'Heading {level}' if level>0 else 'Normal'
            else:
                doc.add_paragraph(text)
            i += 1
            continue

        # table detection
        if is_table_line(line):
            table_block.append(line)
            i += 1
            continue

        # normal paragraph
        if table_block:
            rows = parse_table(table_block)
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Table Grid'
                for r_idx, row in enumerate(rows):
                    for c_idx, cell in enumerate(row):
                        table.rows[r_idx].cells[c_idx].text = cell
            table_block = []
        doc.add_paragraph(line)
        i += 1

    # flush remaining table
    if table_block:
        rows = parse_table(table_block)
        if rows:
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = 'Table Grid'
            for r_idx, row in enumerate(rows):
                for c_idx, cell in enumerate(row):
                    table.rows[r_idx].cells[c_idx].text = cell

    doc.save(docx_path)


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print('Usage: convert_md_to_docx.py input.md output.docx')
        sys.exit(1)
    md_to_docx(sys.argv[1], sys.argv[2])
