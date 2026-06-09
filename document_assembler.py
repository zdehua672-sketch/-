# -*- coding: utf-8 -*-
"""
论文文档组装器 - DocumentAssembler
将文字章节 + 图片一一对应组装成排版好的 DOCX。
支持中文学术论文格式（宋体/黑体、页边距、图注）。

可从 orchestrator 或 paper_writing_agent 独立调用。
"""

import os
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


# ============================================================================
# 1. 中文字体样式工具
# ============================================================================

def _set_run_font(run, font_name='宋体', size=Pt(12), bold=False, color=None):
    """设置 run 的中英文字体"""
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    if color:
        run.font.color.rgb = color
    # 关键：设置中文字体
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def _add_styled_paragraph(doc, text, font_name='宋体', size=Pt(12),
                          bold=False, indent=True, align=None,
                          space_before=0, space_after=Pt(6)):
    """添加带样式的段落"""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before) if isinstance(space_before, (int, float)) else space_before
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    _set_run_font(run, font_name, size, bold)
    return p


# ============================================================================
# 2. DocumentAssembler 主类
# ============================================================================

class DocumentAssembler:
    """
    论文文档组装器。

    使用方式:
        assembler = DocumentAssembler(title='我的论文', paper_type='chinese')
        assembler.add_section('1. 引言', intro_text)
        assembler.add_figure('fig1.png', '图1 季节对比', caption='冬季与春季...')
        assembler.add_section('2. 结果', results_text)
        assembler.add_figure('fig2.png', '图2 相关性')
        assembler.assemble('output.docx')
    """

    PAPER_CONFIGS = {
        'chinese': {
            'body_font': '宋体',
            'heading_font': '黑体',
            'title_font': '黑体',
            'caption_font': '宋体',
            'body_size': Pt(12),
            'heading1_size': Pt(16),
            'heading2_size': Pt(14),
            'heading3_size': Pt(13),
            'caption_size': Pt(10),
            'line_spacing': 1.5,
            'margins': {'top': Cm(2.54), 'bottom': Cm(2.54), 'left': Cm(3.17), 'right': Cm(3.17)},
            'figure_width': Cm(14),  # 图片宽度（单栏）
            'figure_dpi': 300,
        },
        'sci': {
            'body_font': 'Times New Roman',
            'heading_font': 'Arial',
            'title_font': 'Arial',
            'caption_font': 'Times New Roman',
            'body_size': Pt(11),
            'heading1_size': Pt(14),
            'heading2_size': Pt(12),
            'heading3_size': Pt(11),
            'caption_size': Pt(9),
            'line_spacing': 1.5,
            'margins': {'top': Cm(2.54), 'bottom': Cm(2.54), 'left': Cm(2.54), 'right': Cm(2.54)},
            'figure_width': Cm(15),
            'figure_dpi': 600,
        },
        'nature': {
            'body_font': 'Times New Roman',
            'heading_font': 'Arial',
            'title_font': 'Arial',
            'caption_font': 'Times New Roman',
            'body_size': Pt(10),
            'heading1_size': Pt(12),
            'heading2_size': Pt(11),
            'heading3_size': Pt(10),
            'caption_size': Pt(8),
            'line_spacing': 1.5,
            'margins': {'top': Cm(2.54), 'bottom': Cm(2.54), 'left': Cm(2.54), 'right': Cm(2.54)},
            'figure_width': Cm(15),
            'figure_dpi': 600,
        },
    }

    def __init__(self, title=None, paper_type='chinese', language='zh'):
        """
        Parameters
        ----------
        title : str or None, 论文标题
        paper_type : str, 'chinese' / 'sci' / 'nature'
        language : str, 'zh' / 'en'
        """
        self.title = title
        self.paper_type = paper_type
        self.language = language
        self.config = self.PAPER_CONFIGS.get(paper_type, self.PAPER_CONFIGS['chinese'])
        self.doc = Document()
        self._figure_counter = 0
        self._section_counter = 0
        self._figure_map = {}  # {figure_id: (path, caption)}
        self._content_queue = []  # 按顺序存储 (type, data)

        self._setup_document()

    def _setup_document(self):
        """初始化文档样式"""
        doc = self.doc

        # 正文样式
        style = doc.styles['Normal']
        font = style.font
        font.name = self.config['body_font']
        font.size = self.config['body_size']
        style.element.rPr.rFonts.set(qn('w:eastAsia'), self.config['body_font'])
        style.paragraph_format.line_spacing = self.config['line_spacing']
        style.paragraph_format.space_after = Pt(6)

        # 标题样式
        for level in range(1, 4):
            heading_style = doc.styles[f'Heading {level}']
            heading_style.font.name = self.config['heading_font']
            heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), self.config['heading_font'])
            heading_style.font.color.rgb = RGBColor(0, 0, 0)
            size_key = f'heading{level}_size'
            heading_style.font.size = self.config.get(size_key, Pt(12))

        # 页边距
        for section in doc.sections:
            for side, value in self.config['margins'].items():
                setattr(section, f'{side}_margin', value)

        # 添加标题
        if self.title:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(self.title)
            _set_run_font(run, self.config['title_font'], Pt(22), bold=True)
            p.paragraph_format.space_after = Pt(18)

    # ==================== 内容添加接口 ====================

    def add_section(self, heading, text=None, level=1):
        """
        添加章节标题 + 正文。

        Parameters
        ----------
        heading : str, 章节标题
        text : str or None, 正文内容（可后续用 add_text 追加）
        level : int, 标题级别 1/2/3
        """
        self._content_queue.append(('section', {
            'heading': heading, 'text': text, 'level': level
        }))

    def add_text(self, text, bold=False, indent=True, align=None):
        """添加正文段落"""
        self._content_queue.append(('text', {
            'text': text, 'bold': bold, 'indent': indent, 'align': align
        }))

    def add_figure(self, image_path, caption=None, width=None, label=None):
        """
        添加图片 + 图注。图片插入到当前位置。

        Parameters
        ----------
        image_path : str, 图片文件路径 (PNG/JPG/PDF)
        caption : str or None, 图注文字。None 时自动生成 "图X"
        width : Cm or None, 图片宽度。None 时用默认值
        label : str or None, 图片标签（用于交叉引用）
        """
        self._figure_counter += 1
        fig_num = self._figure_counter

        if caption is None:
            if self.language == 'zh':
                caption = f'图{fig_num}'
            else:
                caption = f'Fig. {fig_num}'

        # 确保图注以图号开头
        if self.language == 'zh' and not caption.startswith('图'):
            caption = f'图{fig_num} {caption}'
        elif self.language == 'en' and not caption.startswith('Fig'):
            caption = f'Fig. {fig_num}. {caption}'

        fig_id = label or f'fig_{fig_num}'
        self._figure_map[fig_id] = (image_path, caption)

        self._content_queue.append(('figure', {
            'path': image_path, 'caption': caption,
            'num': fig_num, 'width': width, 'label': fig_id,
        }))

    def add_table(self, table_data, caption=None):
        """
        添加表格。

        Parameters
        ----------
        table_data : list of list, 表格数据（第一行为表头）
        caption : str or None, 表注
        """
        self._content_queue.append(('table', {
            'data': table_data, 'caption': caption,
        }))

    def add_page_break(self):
        """添加分页符"""
        self._content_queue.append(('page_break', {}))

    # ==================== 组装输出 ====================

    def assemble(self, output_path):
        """
        组装并保存 DOCX。

        Parameters
        ----------
        output_path : str, 输出文件路径

        Returns
        -------
        str: 输出文件路径
        """
        for item_type, data in self._content_queue:
            if item_type == 'section':
                self._render_section(data)
            elif item_type == 'text':
                self._render_text(data)
            elif item_type == 'figure':
                self._render_figure(data)
            elif item_type == 'table':
                self._render_table(data)
            elif item_type == 'page_break':
                self.doc.add_page_break()

        # 保存
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        self.doc.save(output_path)
        print(f"[DocumentAssembler] 文档已保存: {output_path}")
        return output_path

    def get_figure_map(self):
        """返回 {figure_id: (path, caption)} 映射"""
        return dict(self._figure_map)

    # ==================== 渲染方法 ====================

    def _render_section(self, data):
        """渲染章节"""
        heading = data['heading']
        level = data.get('level', 1)
        text = data.get('text')

        self.doc.add_heading(heading, level=level)
        if text:
            self._render_text({'text': text, 'bold': False, 'indent': True, 'align': None})

    def _render_text(self, data):
        """渲染正文段落 — 智能合并短句为段落"""
        text = data.get('text', '')
        if not text:
            return

        # 将文本按段落分割
        lines = text.strip().split('\n')
        # 智能合并：连续的普通文本行合并为一个段落
        merged_paragraphs = []
        current_para = []

        for line in lines:
            line = line.strip()
            if not line:
                # 空行 = 段落分隔
                if current_para:
                    merged_paragraphs.append(' '.join(current_para))
                    current_para = []
                continue

            # 标题行不合并
            if line.startswith('#'):
                if current_para:
                    merged_paragraphs.append(' '.join(current_para))
                    current_para = []
                merged_paragraphs.append(('heading', line))
                continue

            # 粗体标题行不合并
            if line.startswith('**') and line.endswith('**'):
                if current_para:
                    merged_paragraphs.append(' '.join(current_para))
                    current_para = []
                merged_paragraphs.append(('bold', line))
                continue

            # 列表项不合并
            if line.startswith('- '):
                if current_para:
                    merged_paragraphs.append(' '.join(current_para))
                    current_para = []
                merged_paragraphs.append(('list', line))
                continue

            # 普通文本 → 合并
            current_para.append(line)

        if current_para:
            merged_paragraphs.append(' '.join(current_para))

        # 渲染合并后的段落
        for item in merged_paragraphs:
            if isinstance(item, tuple):
                tag, para_text = item
                if tag == 'heading':
                    heading_text = para_text.lstrip('#').strip()
                    self.doc.add_heading(heading_text, level=3)
                elif tag == 'bold':
                    _add_styled_paragraph(
                        self.doc, para_text.strip('*'),
                        font_name=self.config['heading_font'],
                        size=Pt(12), bold=True, indent=False
                    )
                elif tag == 'list':
                    _add_styled_paragraph(
                        self.doc, para_text[2:],
                        font_name=self.config['body_font'],
                        size=self.config['body_size'],
                        indent=False,
                        space_before=0, space_after=Pt(2)
                    )
            else:
                # 普通段落
                _add_styled_paragraph(
                    self.doc, item,
                    font_name=self.config['body_font'],
                    size=self.config['body_size'],
                    bold=data.get('bold', False),
                    indent=data.get('indent', True),
                    align=data.get('align'),
                )

    def _render_figure(self, data):
        """渲染图片 + 图注"""
        image_path = data['path']
        caption = data['caption']
        width = data.get('width') or self.config['figure_width']

        # 添加图片
        if os.path.exists(image_path):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(image_path, width=width)
        else:
            # 图片不存在时添加占位符
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f'[图片未找到: {image_path}]')
            _set_run_font(run, self.config['body_font'], Pt(10),
                          color=RGBColor(255, 0, 0))

        # 添加图注
        caption_p = self.doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p.paragraph_format.space_before = Pt(4)
        caption_p.paragraph_format.space_after = Pt(12)
        # 图注：图号加粗，说明文字正常
        fig_num_match = re.match(r'(图\d+|Fig\.\s*\d+\.?)\s*', caption)
        if fig_num_match:
            fig_label = fig_num_match.group(1)
            fig_desc = caption[len(fig_num_match.group(0)):]
            run1 = caption_p.add_run(fig_label + ' ')
            _set_run_font(run1, self.config['caption_font'],
                          self.config['caption_size'], bold=True)
            run2 = caption_p.add_run(fig_desc)
            _set_run_font(run2, self.config['caption_font'],
                          self.config['caption_size'])
        else:
            run = caption_p.add_run(caption)
            _set_run_font(run, self.config['caption_font'],
                          self.config['caption_size'])

    def _render_table(self, data):
        """渲染表格"""
        table_data = data.get('data', [])
        caption = data.get('caption')

        if not table_data:
            return

        # 表注
        if caption:
            caption_p = self.doc.add_paragraph()
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_p.paragraph_format.space_after = Pt(4)
            run = caption_p.add_run(caption)
            _set_run_font(run, self.config['caption_font'],
                          self.config['caption_size'], bold=True)

        # 创建表格
        n_rows = len(table_data)
        n_cols = len(table_data[0]) if table_data else 0
        table = self.doc.add_table(rows=n_rows, cols=n_cols, style='Table Grid')

        for i, row_data in enumerate(table_data):
            for j, cell_text in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = str(cell_text)
                # 表头加粗
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = self.config['caption_size']
                        if i == 0:
                            run.bold = True


# ============================================================================
# 3. 一键组装函数
# ============================================================================

def assemble_paper(sections, figures, output_path, title=None,
                   paper_type='chinese', language='zh'):
    """
    一键组装论文文档。

    Parameters
    ----------
    sections : list of dict
        [{'heading': str, 'text': str, 'level': int}, ...]
    figures : list of dict
        [{'path': str, 'caption': str, 'after_section': int}, ...]
        after_section: 图片插入在第几个章节之后（0-indexed）
    output_path : str, 输出路径
    title : str or None, 论文标题
    paper_type : str, 'chinese'/'sci'/'nature'
    language : str, 'zh'/'en'

    Returns
    -------
    str: 输出路径
    """
    assembler = DocumentAssembler(title=title, paper_type=paper_type, language=language)

    # 按 after_section 分组图片
    fig_by_section = {}
    for fig in figures:
        idx = fig.get('after_section', -1)
        fig_by_section.setdefault(idx, []).append(fig)

    # 组装
    for i, section in enumerate(sections):
        assembler.add_section(
            heading=section['heading'],
            text=section.get('text'),
            level=section.get('level', 1)
        )
        # 在该章节后插入图片
        for fig in fig_by_section.get(i, []):
            assembler.add_figure(
                image_path=fig['path'],
                caption=fig.get('caption'),
                width=fig.get('width'),
            )

    # 末尾图片
    for fig in fig_by_section.get(-1, []):
        assembler.add_figure(
            image_path=fig['path'],
            caption=fig.get('caption'),
            width=fig.get('width'),
        )

    return assembler.assemble(output_path)


def assemble_from_analysis(agent, output_dir=None, paper_type='chinese'):
    """
    从 ScientificAnalysisAgent 的结果直接组装文档。

    Parameters
    ----------
    agent : ScientificAnalysisAgent, 已运行完 run() 的分析 agent
    output_dir : str or None, 输出目录
    paper_type : str, 论文类型

    Returns
    -------
    str: 输出路径
    """
    if output_dir is None:
        output_dir = agent.output_dir

    assembler = DocumentAssembler(
        title='科研数据分析报告',
        paper_type=paper_type,
        language='zh'
    )

    # 分析决策
    assembler.add_section('分析决策', level=1)
    if hasattr(agent, 'orchestrator') and agent.orchestrator:
        for name, decision in agent.orchestrator.recommendations.items():
            status = '✓' if decision.get('do') else '✗'
            assembler.add_text(f'{status} {name}: {decision.get("reason", "")}')

    # 各章节文字 + 对应图表
    fig_dir = output_dir
    fig_files = []
    if os.path.exists(fig_dir):
        fig_files = sorted([
            f for f in os.listdir(fig_dir)
            if f.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff'))
        ])

    if hasattr(agent, 'texts') and agent.texts:
        fig_idx = 0
        for section_name, text in agent.texts.items():
            if not text:
                continue
            assembler.add_section(section_name.replace('_', ' ').title(), level=1)
            assembler.add_text(text)

            # 尝试匹配对应的图
            while fig_idx < len(fig_files):
                fig_path = os.path.join(fig_dir, fig_files[fig_idx])
                fig_name = os.path.splitext(fig_files[fig_idx])[0]
                assembler.add_figure(fig_path, caption=fig_name)
                fig_idx += 1
                break  # 每个章节最多插一张

    # 剩余图片
    if hasattr(agent, 'captions'):
        for fig_type, caption in agent.captions.items():
            if caption:
                assembler.add_text(caption, indent=False)

    output_path = os.path.join(output_dir, 'analysis_report.docx')
    return assembler.assemble(output_path)
